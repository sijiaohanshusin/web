"""Curated first-use editions; instructions remain in helpcenter Markdown."""
import argparse
from io import BytesIO
import json
from pathlib import Path
import re
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from build_help import ASSETS, CONTENT, ROOT, OUT, bookmark, field, inline, link

MANIFEST = Path(__file__).with_name('first_use.json')
QA = ROOT / '.shots/first-use'
AUDIENCES = ('recruit', 'member', 'admin')
# Coordinates refer to the original sanitized screenshots, never fabricated UI.
CROPS = {
    'shared-01-register-channel-choice.png': (85, 280, 1345, 810),
    'recruitment-02-register-step-identity.png': (505, 230, 1230, 760),
    'recruitment-03-register-step-contact.png': (505, 300, 1230, 785),
    'recruitment-08-application-step-interests.png': (960, 410, 1335, 940),
    'recruitment-11-application-step-confirm.png': (960, 615, 1335, 1000),
    'first-use-application-mobile.png': (0, 0, 302, 224),
    'recruitment-12-application-status-submitted.png': (945, 570, 1340, 995),
    'recruitment-07-application-step-direction.png': (100, 742, 825, 925),
    'returning-member-02-register-step-identity.png': (500, 215, 1230, 845),
    'returning-member-03-register-step-contact.png': (505, 300, 1230, 795),
    'shared-02-password-login.png': (505, 330, 1110, 785),
    'workspace-start.png': (275, 212, 930, 566),
    'workspace-card.png': (275, 280, 945, 580),
    'workspace-publish.png': (977, 385, 1425, 678),
    'admin-01-dashboard-overview.png': (0, 0, 232, 465),
    'admin-02-member-search.png': (258, 90, 1100, 560),
    'admin-03-returning-review.png': (1085, 220, 1407, 440),
    'admin-07-application-detail.png': (860, 145, 1410, 530),
    'forum-first-topic-desktop.png': (0, 0, 1331, 220),
    'forum-first-reply.png': (0, 0, 1331, 220),
    'forum-first-topic-mobile.png': (0, 0, 390, 184),
    'forum-queue-review.png': (0, 0, 810, 565),
}


def article(ref):
    if not re.fullmatch(r'(recruit|member|admin)/[a-z-]+', ref):
        raise ValueError(f'Invalid article reference: {ref}')
    _, header, body = (CONTENT / f'{ref}.md').read_text(encoding='utf-8').split('---', 2)
    meta = json.loads(header)
    sections = {}
    heading = ''
    for line in body.splitlines():
        if line.startswith('## '):
            heading = line[3:].strip()
            sections[heading] = []
        elif line.strip() and not line.startswith('![') and not re.match(r'^\*[^*]', line):
            sections.setdefault(heading, []).append(line.strip())
    return meta, sections


def excerpt(ref, section, items=None):
    _, sections = article(ref)
    lines = sections[section]
    if items is not None:
        numbered = {int(m[1]): m[2] for line in lines
                    if (m := re.match(r'^(\d+)\. (.*)', line))}
        return [numbered[index] for index in items]
    return [re.sub(r'^\d+\. ', '', line) for line in lines]


def picture(name):
    with Image.open(ASSETS / name) as source:
        result = source.convert('RGB')
        if name in CROPS:
            box = CROPS[name]
            if not (0 <= box[0] < box[2] <= result.width and 0 <= box[1] < box[3] <= result.height):
                raise ValueError(f'Crop outside source: {name}')
            result = result.crop(box)
        return result


def inventory(config):
    names = list(dict.fromkeys(p[k] for book in AUDIENCES for p in config[book]['pages']
                              for k in ('image', 'image2') if k in p))
    font = ImageFont.truetype('C:/Windows/Fonts/arial.ttf', 13)
    QA.mkdir(parents=True, exist_ok=True)
    for group in range(0, len(names), 9):
        sheet = Image.new('RGB', (1050, 1140), '#eeeeee')
        draw = ImageDraw.Draw(sheet)
        for index, name in enumerate(names[group:group+9]):
            image = picture(name)
            label = f'{group+index+1}. {name}\n{image.width} x {image.height}'
            image.thumbnail((340, 327))
            x, y = index % 3 * 350, index // 3 * 380
            draw.text((x+5, y+5), label, font=font, fill='black')
            sheet.paste(image, (x+5, y+48))
        sheet.save(QA / f'inventory-{group//9+1}.png')


def configure(doc, title, audience):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin, sec.bottom_margin = Mm(17), Mm(17)
    sec.left_margin = sec.right_margin = Mm(19)
    sec.header_distance = sec.footer_distance = Mm(8)
    sec.different_first_page_header_footer = True
    # Word's default template may supply a blue Title border and theme fonts.
    # Remove inherited decoration before applying our explicit typography.
    for element in list(doc.styles.element.iter(qn('w:pBdr'))):
        element.getparent().remove(element)
    for name in ('Normal', 'Title', 'Subtitle', 'Heading 1', 'Heading 2', 'Heading 3', 'Caption'):
        style = doc.styles[name]
        style.font.name = 'Microsoft YaHei'
        style._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        fonts = style._element.get_or_add_rPr().rFonts
        for attr in ('asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme'):
            fonts.attrib.pop(qn('w:' + attr), None)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.italic = False
        style.font.bold = name.startswith('Heading') or name == 'Title'
    normal = doc.styles['Normal']
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)
    snap = OxmlElement('w:snapToGrid'); snap.set(qn('w:val'), '0')
    normal._element.get_or_add_pPr().append(snap)
    for name, size in [('Title', 29), ('Heading 1', 22), ('Heading 2', 12), ('Subtitle', 11)]:
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
    doc.styles['Caption'].font.size = Pt(8.5)
    doc.styles['Caption'].font.color.rgb = RGBColor.from_string('48515B')
    doc.styles['Caption'].paragraph_format.space_after = Pt(5)
    sec.header.paragraphs[0].text = 'HEU ESTA    ' + title + ('    内部使用' if audience == 'admin' else '')
    sec.header.paragraphs[0].style = 'Caption'
    sec.header.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
    p = sec.footer.paragraphs[0]
    p.style = 'Caption'
    p.add_run('首次上手  |  2026.09  |  ')
    field(p, 'PAGE'); p.add_run(' / '); field(p, 'NUMPAGES')
    p.add_run('    '); link(p, '返回目录', '#contents')
    update = OxmlElement('w:updateFields'); update.set(qn('w:val'), 'true')
    doc.settings.element.append(update)
    doc.core_properties.title = title
    doc.core_properties.author = 'HEU ESTA'
    doc.core_properties.subject = '网站首次使用与任务导航'


def add_lines(doc, lines, numbered=True):
    for index, text in enumerate(lines, 1):
        p = doc.add_paragraph()
        if numbered:
            p.paragraph_format.left_indent = Mm(6)
            p.paragraph_format.first_line_indent = Mm(-6)
            p.add_run(f'{index}. ').bold = True
        inline(p, text.removeprefix('- '))


def add_picture(doc, name, label, max_height=85):
    img = picture(name)
    # Show narrow phone captures at a natural readable width, not desktop width.
    width = min(172, 92 if img.width < 600 else 172, max_height * img.width / img.height)
    stream = BytesIO(); img.save(stream, format='PNG'); stream.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(stream, width=Mm(width))
    shape._inline.docPr.set('descr', label)
    doc.add_paragraph(label, 'Caption')
    return {'name': name, 'pixels': list(img.size), 'width_mm': round(width, 1),
            'ppi': round(img.width / (width / 25.4)), 'crop': CROPS.get(name)}


def make_book(config, audience, revision):
    book = config[audience]
    doc = Document(); configure(doc, book['title'], audience)
    evidence = []
    markdown = [f"# {book['title']}", '', book['intro'], '']
    doc.add_paragraph('哈尔滨工程大学电子科技协会', 'Subtitle')
    p = doc.add_paragraph('FIRST USE  /  ' + {'recruit':'01 RECRUIT','member':'02 MEMBER','admin':'03 ADMIN'}[audience])
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(16)
    p.runs[0].font.color.rgb = RGBColor.from_string('126472')
    doc.add_paragraph(book['title'], 'Title')
    doc.add_paragraph(config['edition'], 'Subtitle')
    doc.add_paragraph(book['intro'])
    doc.add_heading('从哪里开始', 2)
    starts = {
        'recruit': [('还没有账号', 1), ('已经注册但未报名', 5), ('已经提交报名', 9)],
        'member': [('需要恢复老会员身份', 1), ('已经审核通过', 3), ('只想学习发帖', 6), ('只想设计公开卡片', 8)],
        'admin': [('第一次接手管理', 1), ('正在处理成员申请', 3), ('需要发布公告', 6), ('已有论坛审核权限', 9)],
    }
    for label, task in starts[audience]:
        p = doc.add_paragraph(); link(p, f'{label}  从任务 {task:02d} 开始', f'#task_{task}')
    doc.add_heading('怎样使用这本手册', 2)
    doc.add_paragraph('按页操作，不必一次做完。每页先读目标，完成步骤后勾选检查点。已有经验时直接点击目录跳转；遇到未覆盖的问题，打开本页末尾的详细帮助。')
    doc.add_heading('适用版本', 2)
    doc.add_paragraph(f'流程核验截至 2026 年 9 月 5 日。网站候选版本 {revision}；常规流程截图核验于 9 月 4 日，成员展示、用户名、资料编辑和论坛新截图核验于 9 月 5 日。每幅图保留原截图编号与日期。')
    doc.add_paragraph('本册对应尚未发布的帮助中心与体验改造版本。正式网站暂未启用的页面，以部署后的最新网页帮助为准。截图仅用于说明操作，不代表演示内容可以直接用于正式业务。', 'Caption')
    p = doc.add_paragraph(); link(p, '打开本册网页帮助', f'https://heuesta.cn/help/{audience}/')
    doc.add_page_break()
    bookmark(doc.add_heading('首次任务目录', 1), 'contents')
    doc.add_paragraph('点击任务名称跳转。也可以打印后逐项完成。')
    for n, page in enumerate(book['pages'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(9)
        p.add_run(f'{n:02d}   ').bold = True
        link(p, page['title'], f'#task_{n}')
        p.add_run('    '); field(p, f'PAGEREF task_{n} \\h')
    doc.add_heading('需要更详细的操作', 2)
    doc.add_paragraph('网页帮助支持搜索、按任务查阅和放大原图。资料上传、项目维护、勋章、职位、安全配置等进阶操作不塞进首次流程，请在网页中查找对应任务。')
    for n, page in enumerate(book['pages'], 1):
        ref = page['ref']; meta, _ = article(ref)
        p = doc.add_paragraph(f'任务 {n:02d} / {len(book["pages"]):02d}', 'Subtitle')
        p.paragraph_format.page_break_before = True
        bookmark(doc.add_heading(page['title'], 1), f'task_{n}')
        markdown.extend([f'## {n:02d} {page["title"]}', '', meta['summary'], ''])
        p = doc.add_paragraph(); p.add_run('目标  ').bold = True; inline(p, meta['summary'])
        if 'lead' in page:
            lead = page['lead']
            lines = excerpt(lead['ref'], lead['section'], lead.get('items'))
            p = doc.add_paragraph(); p.add_run('开始前  ').bold = True; inline(p, ' '.join(lines))
            markdown.extend(lines + [''])
        lines = []
        for selection in page['sections']:
            lines.extend(excerpt(ref, selection[0], selection[1] if len(selection)>1 else None))
        add_lines(doc, lines)
        markdown.extend([f'{i}. {line}' for i, line in enumerate(lines, 1)] + [''])
        image_ref = page.get('image_ref', ref)
        image_meta, _ = article(image_ref)
        for key in ('image','image2'):
            if key not in page:
                continue
            name = page[key]
            if name not in image_meta['screenshots']:
                raise ValueError(f'{image_ref} does not declare {name}')
            number = image_meta['screenshots'].index(name) + 1
            label = f'图 {n:02d}-{1 if key == "image" else 2}  {page["title"]} 操作区'
            image_data = add_picture(doc, name, label, max_height=page.get('image_height', 55 if 'image2' in page else 90))
            image_data.update(article=image_ref, screenshot_number=f'{image_meta["order"]}-{number}', verified=image_meta['verified'])
            evidence.append(image_data)
            markdown.extend([f'![{label}](../../../app/helpcenter/assets/{name})', ''])
        if 'extra' in page:
            extra = page['extra']; lines = excerpt(extra['ref'], extra['section'], extra.get('items'))
            p = doc.add_paragraph(); p.add_run(page.get('extra_label', '接下来') + '  ').bold = True; inline(p, ' '.join(lines))
            markdown.extend(lines + [''])
        p = doc.add_paragraph(); p.add_run('完成检查  ').bold = True
        inline(p, '□ ' + meta['checkpoints'][page.get('check', 0)])
        _, sections = article(ref)
        trouble = next(iter(sections.get('遇到问题', [])), '')
        if trouble and not page.get('omit_trouble'):
            p = doc.add_paragraph(); inline(p, trouble.removeprefix('- '))
            markdown.append(trouble)
        p = doc.add_paragraph(style='Caption')
        p.add_run(f'来源 {ref}  |  核验 {image_meta["verified"]}  |  图源 {image_meta["order"]}-{image_meta["screenshots"].index(page["image"])+1}  ')
        link(p, '详细帮助与原图', f'https://heuesta.cn/help/{ref}/')
        if page.get('finish'):
            p.add_run('  '); link(p, '账号安全与后续任务', 'https://heuesta.cn/help/' + page['finish'] + '/')
        markdown.extend(['', f'[详细帮助](https://heuesta.cn/help/{ref}/)', ''])
    OUT.mkdir(parents=True, exist_ok=True)
    stem = audience + '-first-use'
    doc.save(OUT / f'{stem}.docx')
    (OUT / f'{stem}.md').write_text('\n'.join(markdown), encoding='utf-8')
    return {'book':stem,'source_revision':revision,'expected_pages':len(book['pages'])+2,'images':evidence}


def main():
    parser = argparse.ArgumentParser(); parser.add_argument('--inventory', action='store_true')
    args = parser.parse_args()
    config = json.loads(MANIFEST.read_text(encoding='utf-8'))
    if args.inventory:
        inventory(config); return
    revision = subprocess.check_output(['git','rev-parse','--short','HEAD'], cwd=ROOT, text=True).strip()
    results = [make_book(config, audience, revision) for audience in AUDIENCES]
    (OUT / 'first-use-evidence.json').write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding='utf-8')
    for item in results:
        print(f'{item["book"]}: {item["expected_pages"]} planned pages; {len(item["images"])} screenshots')


if __name__ == '__main__':
    main()
