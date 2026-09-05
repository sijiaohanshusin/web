"""Export task manuals from the exact Markdown sources used by /help/."""
import json
import re
from itertools import count
from pathlib import Path

from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from PIL import Image

ROOT = Path(__file__).resolve().parents[2]
CONTENT = ROOT / "app/helpcenter/content"
ASSETS = ROOT / "app/helpcenter/assets"
OUT = ROOT / "docs/help/dist"
BOOKS = {"recruit": "招新注册手册", "member": "老会员使用手册", "admin": "网站管理手册"}
BOOKMARK_IDS = count(1)


def bookmark(paragraph, name):
    key = str(next(BOOKMARK_IDS))
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), key); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), key)
    paragraph._p.insert(1 if paragraph._p.pPr is not None else 0, start)
    paragraph._p.append(end)


def field(paragraph, code):
    item = OxmlElement("w:fldSimple")
    item.set(qn("w:instr"), code)
    paragraph._p.append(item)


def link(paragraph, label, url):
    element = OxmlElement("w:hyperlink")
    if url.startswith('#'):
        element.set(qn("w:anchor"), url[1:])
    else:
        element.set(qn("r:id"), paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True))
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "147482"); props.append(color)
    run.append(props)
    text = OxmlElement("w:t"); text.text = label; run.append(text); element.append(run)
    paragraph._p.append(element)


def inline(paragraph, text):
    for token in re.split(r"(\*\*.*?\*\*|\[[^\]]+\]\([^)]+\))", text):
        match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if match:
            label, url = match.groups()
            link(paragraph, label, "https://heuesta.cn" + url if url.startswith("/") else url)
        elif token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        else:
            paragraph.add_run(token.replace("`", ""))


def make_book(audience, title):
    items = []
    for path in (CONTENT / audience).glob("*.md"):
        _, header, body = path.read_text(encoding="utf-8").split("---", 2)
        items.append((json.loads(header), body, path.stem))
    items.sort(key=lambda item: item[0]["order"])
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin, sec.bottom_margin = Mm(19), Mm(18)
    sec.left_margin = sec.right_margin = Mm(20)
    for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Caption"):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(0, 0, 0)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)
    snap = OxmlElement('w:snapToGrid'); snap.set(qn('w:val'), '0')
    normal._element.get_or_add_pPr().append(snap)
    for name, size in (("Title", 32), ("Heading 1", 23), ("Heading 2", 15)):
        doc.styles[name].font.size = Pt(size)
        doc.styles[name].paragraph_format.space_before = Pt(12)
        doc.styles[name].paragraph_format.space_after = Pt(7)
    doc.styles["Caption"].font.size = Pt(9)
    doc.styles["Caption"].font.color.rgb = RGBColor.from_string('53616E')
    doc.styles["Subtitle"].font.color.rgb = RGBColor.from_string('147482')
    doc.styles["Subtitle"].font.size = Pt(12)
    doc.styles["Subtitle"].paragraph_format.space_after = Pt(6)
    header = sec.header.paragraphs[0]
    header.text = f"HEU ESTA     {title}" + ("     内部使用" if audience == "admin" else "")
    header.style = "Caption"
    footer = sec.footer.paragraphs[0]
    footer.add_run("2026 招新预发布版    ·    帮助中心离线副本    |    ")
    field(footer, "PAGE")
    footer.add_run(" / ")
    field(footer, "NUMPAGES")
    footer.style = "Caption"
    doc.add_paragraph("哈尔滨工程大学电子科技协会", "Subtitle")
    doc.add_paragraph(title, "Title")
    doc.add_paragraph("按任务查阅的图文操作指南", "Subtitle")
    doc.add_paragraph("不必从第一页读到最后。先从目录找到眼前任务，完成步骤后核对结果。网页版本支持搜索、放大图片和逐项勾选；本文档适合离线查阅与交接。")
    if audience == "admin":
        doc.add_paragraph("内部使用。标注“系统管理员”的任务不向普通站务开放。正式站只作日常必要操作，演示与失败恢复测试应在隔离环境执行。")
    doc.add_paragraph("版本说明：常规流程沿用 2026 年 9 月 4 日正式版截图；成员展示、用户名与批量晋升新规则为 2026 年 9 月 5 日候选版，尚未上线。以每篇核验记录为准。")
    p = doc.add_paragraph(); link(p, "打开网页手册与搜索", f"https://heuesta.cn/help/{audience}/")
    doc.add_paragraph("网页入口随新版部署启用。在部署前，请使用本地验收地址，不把链接尚未启用视为账号故障。")
    doc.add_paragraph("快速使用：目录中的任务名称可跳到本册章节；每章末尾可返回目录。网页版适合边操作边查阅，离线版适合交接与打印。所有截图均为操作说明，不代表任何真实成员的公开授权。")
    doc.add_page_break()
    heading = doc.add_heading("任务目录", 1)
    bookmark(heading, 'tasks_toc')
    doc.add_paragraph("点击任务名称跳转  右侧为本册页码", "Caption")
    for number, (meta, _, slug) in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{number:02d}  ").bold = True
        link(p, meta["title"], f"#task_{slug.replace('-', '_')}")
        if meta["access"] == "admin":
            p.add_run("  系统管理员专用").bold = True
        p.paragraph_format.tab_stops.add_tab_stop(Mm(168), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        p.add_run('\t')
        field(p, f"PAGEREF task_{slug.replace('-', '_')} \\h")
    for number, (meta, body, slug) in enumerate(items, 1):
        doc.add_page_break()
        doc.add_paragraph(f"任务 {number:02d}    {'系统管理员' if meta['access'] == 'admin' else '内部操作' if audience == 'admin' else '操作指南'}", "Subtitle")
        heading = doc.add_heading(re.sub(r"[，。；：、·（）→]", " ", meta["title"]), 1)
        bookmark(heading, f"task_{slug.replace('-', '_')}")
        doc.add_paragraph(meta["summary"])
        p = doc.add_paragraph()
        link(p, "查看网页步骤与可放大截图", f"https://heuesta.cn/help/{audience}/{slug}/")
        figure = 0
        for line in body.strip().splitlines():
            line = line.strip()
            if not line:
                continue
            shot = re.fullmatch(r"!\[([^\]]*)\]\(asset:([^)]+)\)", line)
            if shot:
                label, name = shot.groups()
                path = ASSETS / name
                with Image.open(path) as image:
                    w, h = image.size
                # Keep phone screenshots legible without squeezing long pages to thumbnails.
                width = min(170, 135 * w / h)
                if w < 600:
                    width = min(width, 82)
                p = doc.add_paragraph()
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                shape = p.add_run().add_picture(str(path), width=Mm(width))
                shape._inline.docPr.set('descr', label)
                figure += 1
                caption = doc.add_paragraph(f"图 {number}-{figure}  {label}", "Caption")
                caption.paragraph_format.keep_with_next = False
            elif line.startswith("## "):
                doc.add_heading(re.sub(r"[，。；：、·（）→]", " ", line[3:]), 2)
            elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
                continue
            else:
                p = doc.add_paragraph()
                inline(p, line)
        doc.add_heading("完成检查", 2)
        for check in meta["checkpoints"]:
            doc.add_paragraph("□  " + check)
        doc.add_paragraph(f"核验日期 {meta['verified']}    适用版本 {meta['version']}    截图编号见网页源稿", "Caption")
        p = doc.add_paragraph()
        link(p, "返回任务目录", "#tasks_toc")
    doc.core_properties.title = title
    doc.core_properties.subject = "HEU ESTA 网站任务式操作指南"
    doc.core_properties.author = "HEU ESTA"
    out = OUT / (audience + ".docx")
    doc.save(out)
    print(out)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    for audience, title in BOOKS.items():
        make_book(audience, title)
